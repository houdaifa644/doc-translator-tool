import pytest
from docx import Document
from time import sleep
from stage.translation import traduire_document, mock_reverse             #import the function
from docx.shared import Inches
from io import BytesIO
from PIL import Image
from docx.shared import RGBColor


USE_MOCK = True        # choice to use the mock (faster, no API)



def create_doc(body="", header="", footer=""):  # create a document with optional body, header, and footer
    doc = Document()
    if body:
        doc.add_paragraph(body)
    section = doc.sections[0]
    if header:
        section.header.add_paragraph(header)
    if footer:
        section.footer.add_paragraph(footer)
    return doc


#test for wednesday
def create_combined_test_document():
    doc = Document()

    # 1. Heading styles
    doc.add_heading("Main Title", level=1)
    doc.add_heading("Subsection", level=2)

    # 2. Bold, Underline, Red Color
    para = doc.add_paragraph()
    run = para.add_run("Important text.")
    run.bold = True
    run.underline = True
    run.font.color.rgb = RGBColor(255, 0, 0)  # Red

    # 3. Email content
    doc.add_paragraph("Please contact us at hamza92000@icloud.com for help.")

    # 4. Date content
    doc.add_paragraph("Meeting is scheduled on June 13, 2025.")

    return doc


# Utilities to extract non-empty text
def get_paragraphs(doc):
    return [p.text for p in doc.paragraphs if p.text.strip()]       #get every paragraph not empty  of doc

def get_header(doc):
    return [p.text for p in doc.sections[0].header.paragraphs if p.text.strip()]

def get_footer(doc):
    return [p.text for p in doc.sections[0].footer.paragraphs if p.text.strip()]

# ------------------ TESTS ------------------

def test_document_with_one_line():
    doc = create_doc(body="Hello world")                    #create document with one line code
    translated = traduire_document(doc, use_mock=USE_MOCK)      #traduce if mock is false, while use mock
    result = get_paragraphs(translated)[0]
    assert result != "Hello world"                 #verify the text was changed


def test_document_with_one_paragraph():
    doc = Document()
    para = doc.add_paragraph()
    para.add_run("This is a full paragraph to test the translation.")   #paragraph with one line
    translated = traduire_document(doc, use_mock=USE_MOCK)
    result = get_paragraphs(translated)[0]
    assert "This" not in result                                     #verifiy at least on, so it's OK for the others


def test_document_with_many_paragraphs():
    doc = Document()
    doc.add_paragraph("Paragraph 1.")
    doc.add_paragraph("Paragraph 2.")
    doc.add_paragraph("Paragraph 3.")
    translated = traduire_document(doc, use_mock=USE_MOCK)
    results = get_paragraphs(translated)
    assert len(results) == 3
    assert all("Paragraph" not in r for r in results)          #verify there is 3 para, and none of them has an english word


def test_document_with_table():
    doc = Document()
    table = doc.add_table(rows=1, cols=2)
    cell1, cell2 = table.rows[0].cells
    cell1.text = "Hello"
    cell2.text = "World"

    translated = traduire_document(doc, use_mock=USE_MOCK)

    result1 = translated.tables[0].rows[0].cells[0].text
    result2 = translated.tables[0].rows[0].cells[1].text

    assert result1 != "Hello"
    assert result2 != "World"


def test_document_with_header():
    doc = create_doc(body="Main content", header="This is a header")
    translated = traduire_document(doc, use_mock=USE_MOCK)
    assert get_header(translated)[0] != "This is a header"


def test_document_with_footer():
    doc = create_doc(body="Main content", footer="This is a footer")
    translated = traduire_document(doc, use_mock=USE_MOCK)
    assert get_footer(translated)[0] != "This is a footer"


def test_document_with_header_and_footer():
    doc = create_doc(body="Main content", header="Header goes here", footer="Footer goes here")
    translated = traduire_document(doc, use_mock=USE_MOCK)
    assert get_header(translated)[0] == "ereh seog redaeH"
    assert get_footer(translated)[0] != "Footer goes here"


def test_document_with_bold_line():
    doc = Document()
    para = doc.add_paragraph()
    run = para.add_run("This is a bold line.")
    run.bold = True
    translated = traduire_document(doc, use_mock=USE_MOCK)
    text = get_paragraphs(translated)[0]
    assert "This" not in text


def test_document_with_partially_bold_line():
    doc = Document()
    para = doc.add_paragraph()
    para.add_run("Start normal, ")
    bold_part = para.add_run("bold part")
    bold_part.bold = True
    translated = traduire_document(doc, use_mock=USE_MOCK)
    result = get_paragraphs(translated)[0]
    assert "bold" not in result


def test_document_with_heading():
    doc = Document()
    doc.add_heading("This is a heading", level=1)
    translated = traduire_document(doc, use_mock=USE_MOCK)
    result = get_paragraphs(translated)[0]
    assert "heading" not in result


def test_direction_of_translation():
    if USE_MOCK:
        pytest.skip("Direction check only valid with real Arabic translation")
    
    doc = create_doc(body="Hello everyone")
    translated = traduire_document(doc, use_mock=False)
    translated_text = get_paragraphs(translated)[0]

    assert translated_text != "Hello everyone"
    #assert any("\u0600" <= c <= "\u06FF" for c in translated_text) to verify there is an arabic charac


def test_translation_with_numbers():
    if USE_MOCK:
        pytest.skip("This test requires real API translation (USE_MOCK = False)")

    
    doc = create_doc(body="There are 25 students in the classroom.") # create a document containing a sentence with a number
   
    translated = traduire_document(doc, use_mock=False)  # translate the document using the real translation API
    
    result = get_paragraphs(translated)[0]  # extract the translated text

    print(" translated result:", result)

    assert "25" in result   # the number "25" should still appear in the translated text
    assert "students" not in result #it should be translated



def test_translation_with_bullet_list():
    if USE_MOCK:
        pytest.skip("This test requires real API translation (USE_MOCK = False)")

    # create a document with a bullet list (ListBullet style)
    doc = Document()
    doc.add_paragraph("Apple", style="ListBullet")
    doc.add_paragraph("Banana", style="ListBullet")
    doc.add_paragraph("Orange", style="ListBullet")

    translated = traduire_document(doc, use_mock=False)

    results = get_paragraphs(translated)    #get all para translated

    print(" translated bullet list:", results)

    # check that original English words are not present anymore
    for word in ["Apple", "Banana", "Orange"]:
        assert all(word not in r for r in results)



def test_translation_with_cover_page():
    if USE_MOCK:
        pytest.skip("This test requires real API translation (USE_MOCK = False)")

    # create a simulated cover page with title, subtitle, and date
    doc = Document()
    doc.add_paragraph("Project Title", style="Title")
    doc.add_paragraph("This document describes the objectives.", style="Subtitle")
    doc.add_paragraph("Date: May 29, 2024")

    # translate the document using the real translation API
    translated = traduire_document(doc, use_mock=False)

    # get all non empty translated paragraphs
    results = get_paragraphs(translated)

    print("translated cover page:", results)

    # ensure original English text is no longer present
    assert all("Project Title" not in r for r in results)
    assert all("objectives" not in r for r in results)
    assert all("May 29, 2024" not in r for r in results)




def test_inspect_inline_image_structure():

    image_stream = BytesIO()                                #create virtual file
    image = Image.new("RGB", (100, 100), color="red")       #create a red picture
    image.save(image_stream, format="PNG")                  #save the picture in the virtual file
    image_stream.seek(0)                                    #put back the cursor to the beginning

    doc = Document()                                            #create worddoc
    doc.add_picture(image_stream, width=Inches(2))              #add the picture to the doc 

    shapes = doc.inline_shapes                                      #take picture inline 
    assert len(shapes) == 1                                         # verify there is 1 picture added


    shape = shapes[0]                                                       #data of the picture
    print("Type Python :", type(shape))                                     # the type of the picture (from class docx.shape.InlineShape)
    print(" width x height):", shape.width, "x", shape.height)              #more info


    if hasattr(shape, "image"):                                         #checks shape (inline image) has an .image attribute
        img = shape.image                                               #take the linked image object
        print("name of picture file :", img.filename)
        print("size:", len(img.blob))
        print("type image :", type(img))               
    else:
        print("InlineShape object don't have an image accessible via .image")



def test_image_preserved_after_translation():
    if not USE_MOCK:
        pytest.skip("This test requires mock translation (USE_MOCK = True)")

    image_stream = BytesIO()                                    # create a red image in memory
    image = Image.new("RGB", (100, 100), color="blue")
    image.save(image_stream, format="PNG")
    image_stream.seek(0)

    doc = Document()                                            #create a document and insert the image
    doc.add_picture(image_stream, width=Inches(2))
    doc.add_paragraph("Hello, this is a test document")
    assert len(doc.inline_shapes) == 1                              # initial check: one image expected

    translated = traduire_document(doc, use_mock=True)                  # translate the document

    print(f"Images before: {len(doc.inline_shapes)} | Images after: {len(translated.inline_shapes)}")
    assert len(translated.inline_shapes) == 1                   #Verify that the image was preserved

    translated.save("docs/doc_traduit_avec_image.docx")          #save the file




def test_caption_with_image_translated():
    if not USE_MOCK:
        pytest.skip("This test requires mock translation (USE_MOCK = True)")

    image_stream = BytesIO()
    image = Image.new("RGB", (100, 100), color="green")
    image.save(image_stream, format="PNG")
    image_stream.seek(0)

    # create document with image and caption
    doc = Document()
    doc.add_paragraph().add_run().add_picture(image_stream, width=Inches(2))  # image in its own paragraph
    caption_para = doc.add_paragraph("Figure 1: This is a caption.")         # caption below
    caption_para.style = "Caption"

    #check before translation
    assert len(doc.inline_shapes) == 1
    assert "Figure 1" in get_paragraphs(doc)[0]

    #check caption is an actual para
    # translate
    translated = traduire_document(doc, use_mock=True)

    # check image still there
    assert len(translated.inline_shapes) == 1

    # check caption was translated (inverted by mock)
    translated_texts = get_paragraphs(translated)
    print("→ All paragraphs:", translated_texts)
    print("→ Caption after translation:", translated_texts[0])
    assert "Figure 1" not in translated_texts[0]  # should be inverted by mock

    translated.save("docs/doc_with_image_and_caption.docx")


def test_translation_with_numbers_and_symbols():
    if USE_MOCK:
        pytest.skip("This test requires real API translation (USE_MOCK = False)")

    #create a paragraph with numbers and symbols
    doc = create_doc(body="Order #12345 - Total: 99.99 € ✅")

    translated = traduire_document(doc, use_mock=False)

    result = get_paragraphs(translated)[0]
    print("→ Translated result:", result)

    #check that the number and symbols are still present
    assert "12345" in result
    assert "99.99" in result or "99" in result  # some APIs drop decimals
    assert "€" in result or "EUR" in result


def test_heading_styles_translation():
    if not USE_MOCK:
        pytest.skip("This test requires mock translation (USE_MOCK = True)")

    doc = Document()
    doc.add_heading("Main Title", level=1)          #first heading
    doc.add_heading("Subsection", level=2)          #seconde one

    translated = traduire_document(doc, use_mock=True)

    translated_texts = get_paragraphs(translated)
    print("translated headings:", translated_texts)

    assert "Main Title"[::-1] in translated_texts[0]
    assert translated.paragraphs[0].style.name == "Heading 1"       #ensure that it is still 2 para separated
    assert translated.paragraphs[1].style.name == "Heading 2"


def test_bold_underline_color_preservation():
    if not USE_MOCK:
        pytest.skip("This test requires mock translation (USE_MOCK = True)")

    doc = Document()
    para = doc.add_paragraph()
    run = para.add_run("Important text.")
    run.bold = True                             #bold
    run.underline = True                        #underlined
    run.font.color.rgb = RGBColor(255, 0, 0)  # Red

    translated = traduire_document(doc, use_mock=True)
    t_para = translated.paragraphs[0]
    t_run = t_para.runs[0]

    print("text translated:", t_run.text)

    assert run.text[::-1] == t_run.text
    assert t_run.bold is True
    assert t_run.underline is True
    assert t_run.font.color.rgb == RGBColor(255, 0, 0)


def test_email_preserved():
    if not USE_MOCK:
        pytest.skip("This test requires mock translation (USE_MOCK = True)")

    doc = create_doc(body="Please contact us at hamza92000@icloud.com for help.")
    translated = traduire_document(doc, use_mock=True)

    result = get_paragraphs(translated)[0]
    print("translated:", result)

    assert "hamza92000@icloud.com" in result


def test_date_translation():
    if not USE_MOCK:
        pytest.skip("This test requires mock translation (USE_MOCK = True)")

    doc = create_doc(body="Meeting is scheduled on June 13, 2025.")
    translated = traduire_document(doc, use_mock=True)

    result = get_paragraphs(translated)[0]
    print("translated date:", result)

    assert "5202 ,31 enuJ" in result                    #not enuJ 13, 2025


'''
def test_save_combined_document_for_visual_check():
    if not USE_MOCK:
        pytest.skip("This test requires mock translation (USE_MOCK = True)")

    doc = Document()

    doc.add_heading("Main Title", level=1)
    doc.add_heading("Subsection", level=2)'''

