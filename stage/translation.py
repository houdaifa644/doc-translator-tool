from tenacity import retry, stop_after_attempt, wait_fixed
import requests
from docx import Document
import os
import re
from io import BytesIO
from docx.oxml import parse_xml
from docx.oxml.ns import qn
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.text.paragraph import Paragraph
from docx.table import Table
from docx.shared import Inches
from docx.shape import InlineShape
from docx.enum.text import WD_ALIGN_PARAGRAPH

from stage.utils import traduire_texte  # import the translation function

# this generator yields paragraphs, tables, and images in the same order they appear
def iter_block_items_with_images(doc):
    inline_shapes = list(doc.inline_shapes)
    current_shape_idx = 0

    for block in doc.element.body.iterchildren():
        if isinstance(block, CT_P):
            para = Paragraph(block, doc)
            if "w:drawing" in block.xml and current_shape_idx < len(inline_shapes):
                yield inline_shapes[current_shape_idx]
                current_shape_idx += 1
            yield para
        elif isinstance(block, CT_Tbl):
            yield Table(block, doc)

# fake reverse translation to simulate a translation call
def mock_reverse(text):
    emails = re.findall(r'\S+@\S+\.\S+', text)  # extract all emails
    placeholders = [f"__EMAIL{i}__" for i in range(len(emails))]  # create fake tags

    for email, placeholder in zip(emails, placeholders):
        text = text.replace(email, placeholder)  # protect emails

    reversed_text = text[::-1]  # reverse the text

    for placeholder, email in zip(placeholders, emails):
        reversed_placeholder = placeholder[::-1]  # reverse placeholder back
        reversed_text = reversed_text.replace(reversed_placeholder, email)  # restore email

    return reversed_text

# if translation API fails, retry 3 times with 2 second wait
@retry(stop=stop_after_attempt(3), wait=wait_fixed(2))
def appel_api_libretranslate(texte):
    response = requests.post(
        "https://libretranslate.de/translate",
        data={
            "q": texte,
            "source": "en",
            "target": "ar",
            "format": "text"
        },
        timeout=10
    )
    response.raise_for_status()
    return response.json()["translatedText"]

# main function to translate all docx content
def traduire_document(doc, use_mock=True):
    doc_traduit = Document()
    print(f"there is : {len(doc.inline_shapes)} image(s)")

    for item in iter_block_items_with_images(doc):
        if isinstance(item, Paragraph):
            style_name = item.style.name
            try:
                new_para = doc_traduit.add_paragraph(style=style_name)
            except KeyError:
                new_para = doc_traduit.add_paragraph()  # fallback if style not found

            # copy paragraph spacing and alignment
            new_para.paragraph_format.space_before = item.paragraph_format.space_before
            new_para.paragraph_format.space_after = item.paragraph_format.space_after
            new_para.paragraph_format.left_indent = item.paragraph_format.left_indent
            new_para.paragraph_format.right_indent = item.paragraph_format.right_indent
            new_para.paragraph_format.first_line_indent = item.paragraph_format.first_line_indent
            new_para.alignment = item.alignment

            # if paragraph is fully empty, skip
            if item.text.strip() == "":
                if all(run.text.strip() == "" for run in item.runs):
                    print("> empty paragraph skipped (all runs empty)")
                    continue
                else:
                    print("> empty paragraph copied (some runs not empty)")

            leading_spaces = len(item.text) - len(item.text.lstrip(" "))
            leading_tabs = len(item.text) - len(item.text.lstrip("\t"))
            prefix = " " * leading_spaces + "\t" * leading_tabs

            for i, run in enumerate(item.runs):
                run_text = run.text
                translated_text = traduire_texte(run_text, use_mock)
                run_prefix = prefix if i == 0 else ""
                new_run = new_para.add_run(run_prefix + translated_text)

                # copy basic font style
                new_run.bold = run.bold
                new_run.italic = run.italic
                new_run.underline = run.underline
                new_run.font.name = run.font.name
                new_run.font.size = run.font.size
                new_run.font.color.rgb = run.font.color.rgb if run.font.color and run.font.color.rgb else None
                new_run.font.highlight_color = run.font.highlight_color
                new_run.font.strike = run.font.strike

        elif isinstance(item, Table):
            print("\n=== new table found ===")
            new_table = doc_traduit.add_table(rows=len(item.rows), cols=len(item.columns))

            if item.style:
                new_table.style = item.style
                print(f"> table style : {item.style}")

            for i, row in enumerate(item.rows):
                is_header = (i == 0)
                print(f"--- row {i+1} ---")
                for j, cell in enumerate(row.cells):
                    print(f"  > cell ({i+1},{j+1}) : {cell.text.strip()[:50]}")
                    new_cell = new_table.cell(i, j)
                    new_cell._tc.clear_content()  # remove default content

                    for para_idx, para in enumerate(cell.paragraphs):
                        print(f"    - paragraph {para_idx+1} (alignment: {para.alignment})")
                        new_para = new_cell.add_paragraph()
                        new_para.paragraph_format.space_before = para.paragraph_format.space_before
                        new_para.paragraph_format.space_after = para.paragraph_format.space_after
                        new_para.alignment = para.alignment

                        for run_idx, run in enumerate(para.runs):
                            run_text = run.text
                            translated = traduire_texte(run_text, use_mock)
                            print(f"      • run {run_idx+1}: '{run_text}' ⟶ '{translated}'")
                            new_run = new_para.add_run(translated)

                            # copy run formatting
                            new_run.bold = run.bold
                            new_run.italic = run.italic
                            new_run.underline = run.underline
                            new_run.font.name = run.font.name
                            new_run.font.size = run.font.size
                            if run.font.color and run.font.color.rgb:
                                new_run.font.color.rgb = run.font.color.rgb
                            new_run.font.highlight_color = run.font.highlight_color
                            new_run.font.strike = run.font.strike

                    # if header row, add blue background
                    if is_header:
                        shading = parse_xml(r'''
                            <w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
                                w:val="clear" w:color="auto" w:fill="4F81BD"/>
                        ''')
                        tcPr = new_cell._tc.get_or_add_tcPr()
                        tcPr.append(shading)

                    # add border to all cells
                    tc = new_cell._tc
                    tcPr = tc.get_or_add_tcPr()
                    borders = parse_xml(r'''
                        <w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
                            <w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>
                            <w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>
                            <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>
                            <w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>
                        </w:tcBorders>''')
                    tcPr.append(borders)

        elif isinstance(item, InlineShape):
            # extract image stream from the document
            rId = item._inline.graphic.graphicData.pic.blipFill.blip.embed
            image_part = doc.part.related_parts[rId]
            image_bytes = image_part.blob
            image_stream = BytesIO(image_bytes)

            # add a new paragraph for the image
            para = doc_traduit.add_paragraph()
            para.paragraph_format.space_before = 0
            para.paragraph_format.space_after = 0

            try:
                # try to recover original alignment
                original_para = item._inline.getparent().getparent()
                align_str = original_para.get(qn('w:jc'))

                align_map = {
                    "left": WD_ALIGN_PARAGRAPH.LEFT,
                    "center": WD_ALIGN_PARAGRAPH.CENTER,
                    "right": WD_ALIGN_PARAGRAPH.RIGHT,
                    "both": WD_ALIGN_PARAGRAPH.JUSTIFY
                }

                if align_str in align_map:
                    para.alignment = align_map[align_str]
                    print(f"  inherited alignment : {align_str}")
                else:
                    print("  alignment not found or not valid")

            except Exception as e:
                print(f"  failed to get alignment : {e}")

            # insert the image with original width
            run = para.add_run()
            run.add_picture(image_stream, width=item.width)
            print("  image inserted")

    return doc_traduit