from docx import Document

def traduire_texte_fr_en(texte):
    dico = {
        "bonjour": "hello",
        "et": "and",
        "bienvenue": "welcome",
        "merci": "thank you",
        "pour": "for",
        "votre": "your",
        "lecture": "reading"
    }
    mots = texte.lower().split()
    return ' '.join([dico.get(mot, mot) for mot in mots])


# function to modify header and footer of a word document
def modifier_header_footer(fichier_entree, fichier_sortie, texte_header, texte_footer):
    doc = Document(fichier_entree)  # load the word document

    for section in doc.sections:  # iterate through each section

        # change the header
        if section.header.paragraphs:
            section.header.paragraphs[0].text = texte_header
        else:
            section.header.add_paragraph(texte_header)

        # change the footer
        if section.footer.paragraphs:
            section.footer.paragraphs[0].text = texte_footer
        else:
            section.footer.add_paragraph(texte_footer)

    doc.save(fichier_sortie)  # save the document
    print("header and footer updated and saved")
    
if __name__ == "__main__":
    # chemins relatifs par rapport à stage/
    modifier_header_footer("../docs/doc_complexe.docx", "../docs/doc_modifie.docx")
